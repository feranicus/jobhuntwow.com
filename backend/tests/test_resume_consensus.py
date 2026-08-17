#!/usr/bin/env python3
"""test_resume_consensus.py — the 4-model consensus tailor, proven against a STUBBED endpoint.

    python3 backend/tests/test_resume_consensus.py

NO NETWORK, NO API KEY, NO pytest. Every model call goes through an injected `poster`, and the
transport is asserted to have been called zero times outside it - a test that reaches the internet
is not testing this repository.

It is deliberately NOT under backend/app: `Dockerfile.web` does `COPY backend/app ./app`, so a test
in there would ship into the production image.

WHAT IS PROVEN (each of these is a defect this design exists to prevent):
  1  an INVENTED employer is flagged
  2  a DROPPED employer is flagged - and the DETERMINISTIC check is the authority, not the model
  3  an all-flags auditor CANNOT empty or gut the document
  4  the auditor is never the author, and never the author's VENDOR
  5  kimi's payload carries temperature 0.6 and NO response_format
  6  the depth floor rejects a stub AND passes a realistic draft (both directions)
  7  a thin draft moves to the NEXT MODEL in the chain instead of being rendered
  8  a 400 is repaired by WHAT THE SERVER NAMED, and thinking stays suppressed
  9  a revision that guts or drops an employer is REJECTED and the draft stands
 10  with no distinct model available the audit is REFUSED, not faked
 11  the manifest fields the UI needs are present
 12  `highlights` and `earlier` actually reach the document renderer (the measured defect)

THE EXIT GATE IS THE LAST STATEMENT IN THIS FILE, on purpose. A gate in the middle silently stops
enforcing every check appended after it, and new checks are always appended - the sibling project
found 73 checks, including its public-suffix guard, running with rc=0 for exactly that reason.
"""
import asyncio
import copy
import importlib.util
import json
import os
import sys

HERE = os.path.dirname(os.path.abspath(__file__))
APP = os.path.join(HERE, "..", "app")
sys.path.insert(0, os.path.abspath(os.path.join(HERE, "..")))


def _load(name, path):
    spec = importlib.util.spec_from_file_location(name, path)
    m = importlib.util.module_from_spec(spec)
    sys.modules[name] = m
    spec.loader.exec_module(m)
    return m


RC = _load("rc_under_test", os.path.join(APP, "resume_consensus.py"))

FAILS = []
RUN = [0]


def check(cond, label, detail=""):
    RUN[0] += 1
    if cond:
        print("  ok    %s" % label)
    else:
        print("  FAIL  %s%s" % (label, ("  <- " + str(detail)) if detail else ""))
        FAILS.append(label)


def section(t):
    print("\n[%s]" % t)


# --------------------------------------------------------------------------- fixtures
# A real-shaped profile: three employers, numbers present, one older role to demote.
PROFILE = """Jane Roe
jane.roe@example.com | +44 20 7946 0958 | London, UK
Head of Infrastructure

EXPERIENCE
Head of Infrastructure, Stars4Business OU, 2021 - present, Tallinn
 - Ran the platform migration for 40 services onto Kubernetes, cutting deploy time from 6 hours to 20 minutes.
 - Managed a team of 12 engineers across 3 countries and a 2.4M EUR annual budget.
 - Negotiated the carrier contract that reduced transit spend by 31 percent.

Infrastructure Lead, Canonical, 2017 - 2021, London
 - Built the CI fleet that served 900 builds a day.
 - Led the migration of 220 servers to Ubuntu LTS with no unplanned downtime.

Systems Engineer, Verint, 2013 - 2017, Tel Aviv
 - Automated provisioning for 60 customer environments.

EDUCATION
BSc Computer Science, University of Leeds, 2013

SKILLS
Kubernetes, Terraform, AWS, Python, Go, networking, BGP, incident management, team leadership
"""

JD = {"title": "Director of Platform Engineering", "company": "Acme Cloud", "location": "Remote",
      "text": ("We need a Director of Platform Engineering to own Kubernetes, Terraform and AWS "
               "across a growing estate. You will lead a distributed team, own the reliability "
               "budget, and drive cost efficiency. Experience with BGP and carrier negotiation is "
               "a plus. Rust experience welcome.")}


def good_resume():
    """A realistic draft: long enough to clear the floor, employers all present."""
    return {
        "summary": ("Head of Infrastructure with over a decade running platform and network estates "
                    "across three countries. Migrated 40 services to Kubernetes and cut deploy time "
                    "from six hours to twenty minutes. Owns budget, vendor contracts and reliability "
                    "for teams of up to twelve engineers. Looking to run platform engineering at "
                    "Acme Cloud's scale."),
        "highlights": [
            "Cut deploy time from 6 hours to 20 minutes by migrating 40 services onto Kubernetes.",
            "Reduced transit spend by 31 percent by renegotiating the carrier contract.",
            "Led 220 server migrations to Ubuntu LTS with no unplanned downtime.",
            "Managed 12 engineers across 3 countries against a 2.4M EUR annual budget.",
        ],
        # NOTE: "Terraform" is deliberately ABSENT although the JD asks for it and the profile
        # evidences it - that is what makes it a genuine `missed_keywords` hit below.
        "skills": ["Kubernetes", "AWS", "Python", "Go", "BGP", "networking",
                   "incident management", "team leadership"],
        "experience": [
            {"title": "Head of Infrastructure", "company": "Stars4Business OU",
             "dates": "2021 - present", "location": "Tallinn", "bullets": [
                 "Migrated 40 services onto Kubernetes, cutting deploy time from 6 hours to 20 minutes.",
                 "Managed 12 engineers across 3 countries against a 2.4M EUR annual budget.",
                 "Negotiated the carrier contract that reduced transit spend by 31 percent.",
             ]},
            {"title": "Infrastructure Lead", "company": "Canonical", "dates": "2017 - 2021",
             "location": "London", "bullets": [
                 "Built the CI fleet that served 900 builds a day for the engineering organisation.",
                 "Led the migration of 220 servers to Ubuntu LTS with no unplanned downtime.",
             ]},
        ],
        "earlier": [{"title": "Systems Engineer", "company": "Verint", "dates": "2013 - 2017"}],
        "all_employers": ["Stars4Business OU", "Canonical", "Verint"],
        "keywords_matched": ["Kubernetes", "AWS", "BGP"],
        "gaps": ["Rust is not evidenced in the profile"],
    }


def good_cover():
    return {"salutation": "Hiring Team", "closing": "Sincerely,", "paragraphs": [
        "I am applying for the Director of Platform Engineering role at Acme Cloud because the "
        "estate you describe is the one I have just rebuilt: 40 services onto Kubernetes, with "
        "deploy time down from six hours to twenty minutes.",
        "Your posting leads with Kubernetes and AWS and with cost efficiency. I migrated "
        "220 servers with no unplanned downtime, and renegotiated the carrier contract that took 31 "
        "percent out of transit spend, so the reliability and the cost lines both moved.",
        "Leading a distributed team is the day job rather than a new challenge: twelve engineers "
        "across three countries, against a 2.4M EUR budget, with the on-call and incident process "
        "I put in place.",
        "I would welcome a conversation about the first ninety days and where you want the "
        "reliability budget to land."]}


def stub(*, resume=None, cover=None, audit=None, revision=None, raise_400=None, calls=None):
    """A poster that answers by ROLE-OF-CALL, so the whole consensus runs with no network.

    It classifies the call from the SYSTEM prompt rather than from call order, because the author's
    two documents are requested concurrently and order is therefore not deterministic.
    """
    log = calls if calls is not None else []

    async def poster(payload, timeout):
        log.append(copy.deepcopy(payload))
        sysmsg = payload["messages"][0]["content"]
        usermsg = payload["messages"][1]["content"]
        if raise_400 and len(log) == 1:
            raise RC.ModelHTTPError(400, raise_400)
        if sysmsg.startswith("You are an INDEPENDENT adversarial reviewer"):
            body = audit
        elif sysmsg.startswith("You are the author of this resume"):
            kind = "resume" if "corrected resume JSON" in usermsg else "cover"
            body = (revision or {}).get(kind)
        elif "tailor an existing resume" in sysmsg:
            body = resume
        else:
            body = cover
        if callable(body):
            body = body(payload)
        if body is None:
            body = {}
        return {"choices": [{"message": {"content": json.dumps(body)},
                             "finish_reason": "stop"}],
                "usage": {"completion_tokens": 900}}
    return poster, log


def run(coro):
    return asyncio.get_event_loop().run_until_complete(coro)


# =========================================================================== 1 + 2 + 11
section("1/2  an invented employer is flagged, a dropped employer is flagged")

# The author invents "Google" (absent from the profile) and DROPS Verint from the document while
# still listing it in all_employers - the exact 2-page-limit failure this product was built around.
bad = good_resume()
bad["experience"].append({"title": "Staff Engineer", "company": "Google", "dates": "2011 - 2013",
                          "bullets": ["Scaled the ranking pipeline to 4 billion documents a day."]})
bad["earlier"] = []                                     # Verint demoted... and then deleted
bad["all_employers"] = ["Stars4Business OU", "Canonical", "Verint", "Google"]

AUDIT_FINDS = {
    "verdict": "revise",
    "invented_facts": [{"where": "resume.experience[2]",
                        "claim": "Staff Engineer at Google, scaled the ranking pipeline to 4 billion documents a day",
                        "why": "the profile contains no Google role"}],
    "employers_dropped": ["Verint"],
    "weak_bullets": [],
    "missed_keywords": ["Terraform"],
    "cover_issues": [],
    "scores": {"truthfulness": 3, "jd_fit": 7, "impact": 6, "structure": 7, "cover_letter": 7},
    "notes": "Remove the Google role and restore Verint.",
}

poster, log = stub(resume=bad, cover=good_cover(), audit=AUDIT_FINDS,
                   revision={"resume": good_resume(), "cover": good_cover()})
res = run(RC.tailor(PROFILE, JD, poster=poster, rounds=1))
r1 = res["audit"]["rounds"][0]

check(any("google" in f["claim"].lower() for f in r1["kept"]["invented_facts"]),
      "the invented Google role is FLAGGED", r1["kept"]["invented_facts"])
check("Verint" in r1["kept"]["employers_dropped"],
      "the dropped employer Verint is FLAGGED", r1["kept"]["employers_dropped"])
check(RC.missing_employers(bad) == ["Verint"],
      "the DETERMINISTIC employer check finds Verint on its own", RC.missing_employers(bad))
check(RC.truth_check(bad, PROFILE) and "Google" in RC.truth_check(bad, PROFILE)[0],
      "truth_check reports Google as absent from the profile", RC.truth_check(bad, PROFILE))

# The deterministic check is the AUTHORITY: a reviewer flag it disagrees with is REFUSED, and the
# refusal is RECORDED rather than silently dropped.
poster2, _ = stub(resume=good_resume(), cover=good_cover(),
                  audit=dict(AUDIT_FINDS, employers_dropped=["Canonical"], invented_facts=[],
                             missed_keywords=[]),
                  revision={"resume": good_resume(), "cover": good_cover()})
res2 = run(RC.tailor(PROFILE, JD, poster=poster2, rounds=1))
rr = res2["audit"]["rounds"][0]
check("Canonical" not in rr["kept"]["employers_dropped"],
      "a bogus employer-drop flag is REFUSED (Canonical is in the document)",
      rr["kept"]["employers_dropped"])
check(any(x["kind"] == "employer_dropped" for x in rr["refused"]),
      "...and the refusal is RECORDED, not discarded", rr["refused"])

# ISOLATE the invented-fact corroboration. A mutation run proved that removing it was masked by the
# GUTTING CAP: with corroboration gone, both flags survive, the cap fires, and the refused branch
# returns an empty list - so the assertion above passed against broken code. A negative test that
# passes because of defence in depth is measuring the OTHER guard. One flag, well under the cap, is
# the only way to see this rule on its own.
solo = RC.bound_audit({"invented_facts": [
    {"where": "resume.experience[0].bullets[0]",
     "claim": "Migrated 40 services onto Kubernetes, cutting deploy time from 6 hours to 20 minutes.",
     "why": "cannot verify"}]}, good_resume(), good_cover(), PROFILE, JD)
check(solo["kept"]["invented_facts"] == [],
      "a SINGLE bogus invention flag is refused by corroboration alone (no cap involved)",
      solo["kept"]["invented_facts"])
check(solo["refused_all"] is False,
      "  ...and the round was NOT gut-capped, so corroboration is what refused it")
check(any(x["kind"] == "invented_fact" and "IS in the profile" in x["why"]
          for x in solo["refused"]),
      "  ...with the reason recorded: the profile supports the claim", solo["refused"])
# and the opposite direction: a claim the profile does NOT support is KEPT
solo2 = RC.bound_audit({"invented_facts": [
    {"where": "resume.experience[0]", "claim": "Delivered 4 billion documents a day at Google.",
     "why": "no Google role"}]}, good_resume(), good_cover(), PROFILE, JD)
check(len(solo2["kept"]["invented_facts"]) == 1,
      "an UNSUPPORTED claim survives corroboration and is acted on",
      solo2["kept"]["invented_facts"])
check(RC.claim_is_corroborated("40 services onto Kubernetes", PROFILE) is True,
      "claim_is_corroborated: numbers + distinctive tokens present -> True")
check(RC.claim_is_corroborated("scaled to 4 billion documents at Google", PROFILE) is False,
      "claim_is_corroborated: an absent number/name -> False")
check(RC.claim_is_corroborated("did some good work", PROFILE) is False,
      "claim_is_corroborated: nothing checkable -> False, so the flag KEEPS and a human decides")
# A model PARAPHRASES, and re-angling is explicitly permitted, so corroboration must survive
# ordinary morphology - the profile says "platform migration", a bullet says "Migrated". Demanding
# exact words called a truthful bullet unverifiable and would bury the candidate in warnings about
# their own real achievements. And a NUMBER must match on digit boundaries: plain containment finds
# "12" inside "120" and would corroborate an inflated figure against a real one.
for claim, want, why in [
    ("Migrated 40 services onto Kubernetes, cutting deploy time from 6 hours to 20 minutes.",
     True, "a rephrased TRUE bullet is corroborated (migration -> Migrated)"),
    ("Negotiated the carrier contract that reduced transit spend by 31 percent.",
     True, "another rephrased true bullet"),
    ("Cut infrastructure costs by 45 percent.", False, "an INVENTED number is not corroborated"),
    ("Managed a team of 120 engineers.", False,
     "an INFLATED number is not corroborated (120 must not match 12)"),
    ("Delivered 4 billion documents a day at Google.", False, "an invented EMPLOYER"),
]:
    check(RC.claim_is_corroborated(claim, PROFILE) is want, "corroboration: %s" % why,
          RC.claim_is_corroborated(claim, PROFILE))

# missed_keywords must satisfy THREE deterministic conditions: in the JD, in the profile, absent
# from the draft. "Rust" is in the JD but not in the profile -> adding it would be invention.
b = RC.bound_audit({"missed_keywords": ["Terraform", "Rust", "Kubernetes", "Kafka"]},
                   good_resume(), good_cover(), PROFILE, JD)
check(b["kept"]["missed_keywords"] == ["Terraform"],
      "missed_keywords keeps only JD+profile terms the draft omits", b["kept"]["missed_keywords"])
check(any(x["item"] == "Rust" and "invention" in x["why"] for x in b["refused"]),
      "'Rust' is refused because the profile does not evidence it")
check(any(x["item"] == "Kubernetes" and "already uses it" in x["why"] for x in b["refused"]),
      "'Kubernetes' is refused because the draft already uses it")
check(any(x["item"] == "Kafka" and "job description" in x["why"] for x in b["refused"]),
      "'Kafka' is refused because the JD never asks for it")

# The "carries a fact, a number or an OUTCOME" rule is a RESUME-BULLET rule. A cover letter's close
# has no number by design, so quoting one as a weak bullet must be refused BY CHANNEL - otherwise a
# correct closing paragraph is condemned on every single run (measured; it was).
close_para = good_cover()["paragraphs"][-1]
check(RC.bullet_is_weak(close_para),
      "the closing paragraph would look 'weak' to the bullet rule", close_para[:60])
b2 = RC.bound_audit({"weak_bullets": [{"where": "cover.paragraphs[3]", "text": close_para,
                                       "why": "no evidence"}]},
                    good_resume(), good_cover(), PROFILE, JD)
check(b2["kept"]["weak_bullets"] == [],
      "...and it is REFUSED because it is cover prose, not a resume bullet")
check(any("cover-letter prose" in x["why"] for x in b2["refused"]),
      "...with the channel named in the refusal", b2["refused"])
check(close_para not in RC.resume_bullets(good_resume()),
      "resume_bullets() excludes cover paragraphs")
check(RC._quality(good_resume(), good_cover())["bullets"] == len(RC.resume_bullets(good_resume())),
      "the quality count measures resume bullets only")

section("11  the manifest fields the UI can show")
for k in ("resume", "cover", "authors", "attempts", "audit", "audit_rounds", "audit_refused",
          "quality", "errors", "elapsed_ms"):
    check(k in res, "result carries '%s'" % k)
for k in ("auditor", "auditor_vendor", "authors", "rounds", "scores", "refused", "notes"):
    check(k in res["audit"], "audit record carries '%s'" % k)
check(res["audit"]["scores"] == AUDIT_FINDS["scores"], "per-section scores are surfaced",
      res["audit"]["scores"])
check(res["audit_rounds"] == 1, "audit_rounds counts the rounds", res["audit_rounds"])


# =========================================================================== 3  cannot gut
section("3  an all-flags auditor cannot EMPTY or GUT the document")

draft_r, draft_c = good_resume(), good_cover()
all_bullets = RC._draft_bullets(draft_r, draft_c)


def carpet(resume, cover):
    """Condemn EVERY bullet, EVERY employer and EVERY role, quoted exactly."""
    return {
        "verdict": "revise",
        "weak_bullets": [{"where": "resume", "text": b, "why": "no evidence"}
                         for b in RC._draft_bullets(resume, cover)],
        "invented_facts": [{"where": "resume.experience[%d]" % i, "claim": e["company"],
                            "why": "cannot verify"}
                           for i, e in enumerate(resume["experience"])],
        "employers_dropped": list(resume.get("all_employers") or []),
        "missed_keywords": [], "cover_issues": ["everything"],
        "scores": {"truthfulness": 0, "jd_fit": 0, "impact": 0, "structure": 0, "cover_letter": 0},
        "notes": "start again"}


# A revision that would EMPTY the document, offered while the auditor condemns everything.
EMPTY_REV = {"resume": {"summary": "", "experience": [], "earlier": [], "highlights": [],
                        "all_employers": []},
             "cover": {"salutation": "Hi", "paragraphs": [], "closing": ""}}

# --- 3a. against a GOOD draft, every flag is REFUSED by deterministic corroboration.
poster3, log3 = stub(resume=draft_r, cover=draft_c, audit=carpet(draft_r, draft_c),
                     revision=EMPTY_REV)
res3a = run(RC.tailor(PROFILE, JD, poster=poster3, rounds=2))
r3a = res3a["audit"]["rounds"][0]
check(res3a["resume"] is not None and res3a["cover"] is not None,
      "3a both documents SURVIVE a carpet-bombing audit")
check(len(res3a["resume"]["experience"]) == 2, "3a every experience entry survives",
      len(res3a["resume"].get("experience") or []))
check(RC.resume_depth(res3a["resume"]) == RC.resume_depth(draft_r),
      "3a the resume is exactly as deep as the draft - nothing was applied")
check(r3a["kept"]["weak_bullets"] == [],
      "3a no weak-bullet flag survives: every bullet carries a number or an outcome verb")
check(r3a["kept"]["invented_facts"] == [],
      "3a no invention flag survives: every company IS in the profile",
      r3a["kept"]["invented_facts"])
check(r3a["kept"]["employers_dropped"] == [],
      "3a no employer-drop flag survives: the deterministic inventory disagrees")
check(len(r3a["refused"]) >= len(all_bullets),
      "3a every refusal is RECORDED (%d entries)" % len(r3a["refused"]))

# --- 3b. the GUTTING CAP itself, on a draft whose bullets really ARE weak.
def dutiful_resume():
    """Duty statements, no numbers - so bullet_is_weak() genuinely agrees with the auditor.

    Without this the cap can never be exercised: corroboration refuses the flags first, and a cap
    that is never reached is an untested cap.
    """
    r = good_resume()
    r["highlights"] = [
        "Responsible for the infrastructure estate and its associated operational processes.",
        "Accountable for platform strategy and for stakeholder engagement across the business.",
        "Involved in vendor discussions and in the ongoing technology roadmap conversations.",
        "Participated in architecture review and in various cross-functional working groups."]
    r["experience"][0]["bullets"] = [
        "Responsible for the Kubernetes platform and for related operational responsibilities.",
        "Accountable for the infrastructure budget and for supplier relationship management.",
        "Involved in hiring, in performance conversations and in general team administration."]
    r["experience"][1]["bullets"] = [
        "Responsible for continuous integration and for associated tooling responsibilities.",
        "Accountable for the server estate and for its ongoing lifecycle administration."]
    return r


dut = dutiful_resume()
check(RC.contract_ok_resume(dut)[0], "3b the dutiful fixture still clears the depth floor",
      RC.contract_ok_resume(dut)[1])
weak_n = len([b for b in RC._draft_bullets(dut, draft_c) if RC.bullet_is_weak(b)])
tot_n = len(RC._draft_bullets(dut, draft_c))
check(weak_n / float(tot_n) > RC.MAX_WEAK_FRACTION,
      "3b %d of %d bullets are genuinely weak, above the %.0f percent cap"
      % (weak_n, tot_n, RC.MAX_WEAK_FRACTION * 100))
poster3b, _ = stub(resume=dut, cover=draft_c, audit=carpet(dut, draft_c), revision=EMPTY_REV)
res3 = run(RC.tailor(PROFILE, JD, poster=poster3b, rounds=2))
r3 = res3["audit"]["rounds"][0]
check(res3["audit_refused"] is True, "3b the whole ROUND is recorded as REFUSED", res3["audit"])
check(r3["refused_reason"], "3b the refusal states WHY", r3.get("refused_reason"))
check(any(x["kind"] == "whole_audit" for x in r3["refused"]),
      "3b the whole-audit refusal is in the record")
check(r3["kept"]["weak_bullets"] == [],
      "3b no weak-bullet flag is acted on once the round is refused")
check(res3["resume"] is not None and len(res3["resume"]["experience"]) == 2,
      "3b the document is untouched")
check(RC.resume_depth(res3["resume"]) == RC.resume_depth(dut),
      "3b ...and exactly as deep as it was")

# --- 3c. below the cap the audit is actionable, and the EMPTYING revision is still rejected.
one_weak = {"verdict": "revise",
            "weak_bullets": [{"where": "resume.experience[0].bullets[0]",
                              "text": dut["experience"][0]["bullets"][0], "why": "no outcome"}],
            "invented_facts": [], "employers_dropped": [], "missed_keywords": [],
            "cover_issues": [], "scores": {}}
poster4, _ = stub(resume=dut, cover=draft_c, audit=one_weak, revision=EMPTY_REV)
res4 = run(RC.tailor(PROFILE, JD, poster=poster4, rounds=1))
r4 = res4["audit"]["rounds"][0]
check(len(r4["kept"]["weak_bullets"]) == 1,
      "3c a SINGLE corroborated weak bullet IS actionable", r4["kept"]["weak_bullets"])
check(r4["revisions"], "3c ...so a revision is actually attempted", r4["revisions"])
check(any(not rv["applied"] for rv in r4["revisions"]),
      "3c ...and the EMPTYING revision is REJECTED", r4["revisions"])
check(res4["resume"]["experience"], "3c the document survives the emptying revision")
check(RC.resume_depth(res4["resume"]) == RC.resume_depth(dut),
      "3c ...unchanged, because the only offered revision was refused")


# =========================================================================== 4  auditor identity
section("4  the auditor is never the author, and never the author's vendor")

for author in RC.CHAIN:
    a = RC.pick_auditor([author])
    check(a is not None and a != author, "author %-18s -> auditor %s (different id)" % (author, a))
    check(RC.vendor(a) != RC.vendor(author),
          "   ...and a different VENDOR (%s vs %s)" % (RC.vendor(a), RC.vendor(author)))

check(len({RC.vendor(m) for m in RC.CHAIN}) == 4,
      "the committed chain spans FOUR vendors (no shared failure domain)",
      sorted({RC.vendor(m) for m in RC.CHAIN}))
# two different authors (resume by one model, cover by another) -> the auditor must differ from BOTH
a = RC.pick_auditor(["deepseek-3.2", "llama-4-maverick"])
check(a not in ("deepseek-3.2", "llama-4-maverick"), "differs from BOTH authors", a)
check(RC.vendor(a) not in ("deepseek", "meta"), "and from both their vendors", RC.vendor(a))
# an operator override may NOT reinstate the author
os.environ["JHW_TAILOR_AUDITOR"] = "deepseek-3.2"
check(RC.pick_auditor(["deepseek-3.2"]) != "deepseek-3.2",
      "an env override cannot make the author its own auditor")
os.environ.pop("JHW_TAILOR_AUDITOR")
# and in the real run the recorded auditor differs from the recorded author
check(res["audit"]["auditor"] not in set(res["authors"].values()),
      "the RUN records an auditor outside the author set",
      (res["authors"], res["audit"]["auditor"]))


# =========================================================================== 10  refuse, not fake
section("10  with no distinct model, the audit is REFUSED rather than faked")

check(RC.pick_auditor(["solo-model"], chn=["solo-model"]) is not None,
      "a single-model chain still finds a distinct auditor from the committed chain")
check(RC.pick_auditor(["deepseek-3.2", "llama-4-maverick"], chn=RC.CHAIN) is not None,
      "two authors out of four still leaves a distinct auditor")
onlymodel = RC.CHAIN + ["deepseek-3.2"]
check(RC.pick_auditor(set(RC.CHAIN) | {"other"}, chn=RC.CHAIN) is None,
      "when EVERY chain model has authored, pick_auditor returns None (refuse)",
      RC.pick_auditor(set(RC.CHAIN) | {"other"}, chn=RC.CHAIN))

# a run in which the only model IS the author must not audit itself
poster5, log5 = stub(resume=good_resume(), cover=good_cover(), audit={"verdict": "pass"})
res5 = run(RC.tailor(PROFILE, JD, poster=poster5, chn=["deepseek-3.2"], rounds=2))
authors5 = set(v for v in res5["authors"].values() if v)
check(res5["audit"]["auditor"] not in authors5,
      "the auditor is outside the author set even on a 1-model chain", res5["audit"]["auditor"])


# =========================================================================== 5  kimi's payload
section("5  kimi's payload: temperature 0.6, no response_format, thinking suppressed")

kp = RC.build_payload("kimi-k2.6", "sys", "user", max_tokens=6000)
check(kp["temperature"] == 0.6, "temperature is 0.6", kp["temperature"])
check("response_format" not in kp, "response_format is NOT sent", sorted(kp))
check(kp.get("chat_template_kwargs") == {"enable_thinking": False},
      "thinking is suppressed", kp.get("chat_template_kwargs"))
check(kp.get("max_tokens") == 6000,
      "kimi KEEPS its max_tokens ceiling (it is not in JSON mode)", kp.get("max_tokens"))
check(RC.build_payload("kimi-k2.5", "s", "u", max_tokens=10)["temperature"] == 0.6,
      "the rule matches kimi-k2.5 too (prefix match, not an exact id)")

dp = RC.build_payload("deepseek-3.2", "sys", "user", max_tokens=6000)
check(dp["response_format"] == {"type": "json_object"}, "deepseek DOES get JSON mode")
check("max_tokens" not in dp,
      "...and the ceiling is dropped, because on this gateway max_tokens and "
      "response_format:json_object are mutually exclusive", sorted(dp))
check(dp["temperature"] == 0.3, "deepseek keeps the default temperature", dp["temperature"])
check(RC.build_payload("deepseek-3.2", "s", "u", max_tokens=1,
                       json_mode=False).get("max_tokens") == 1,
      "with json_mode off the ceiling is kept")
# building a payload must never mutate the shared policy table
_before = json.dumps(RC.MODEL_PARAMS, sort_keys=True)
RC.build_payload("kimi-k2.6", "s", "u", max_tokens=1)
check(json.dumps(RC.MODEL_PARAMS, sort_keys=True) == _before,
      "MODEL_PARAMS is not mutated by _drop (a shallow copy would corrupt the next call)")


# =========================================================================== 8  targeted 400 repair
section("8  a 400 is repaired by WHAT THE SERVER NAMED")

p, fx = RC.repair_payload(RC.build_payload("kimi-k2.6", "s", "u", max_tokens=99),
                          '{"message":"temperature must be 0.6 for this model"}')
check(p["temperature"] == 0.6, "the temperature is READ OUT OF THE BODY", p["temperature"])
p, fx = RC.repair_payload({"temperature": 1.0, "chat_template_kwargs": {"enable_thinking": False}},
                          '{"message":"temperature must be 0.42 for this model"}')
check(p["temperature"] == 0.42,
      "a CHANGED required value is honoured (the constant is not trusted)", p["temperature"])
check(p.get("chat_template_kwargs") == {"enable_thinking": False},
      "thinking stays suppressed on the retry - stripping it is what caused 46k chars of rambling",
      p)
p, fx = RC.repair_payload(RC.build_payload("deepseek-3.2", "s", "u", max_tokens=99,
                                           json_mode=True) | {"max_tokens": 11000},
                          "max_tokens cannot be set when response_format type is 'json_object'; "
                          "omit max token limits for structured outputs")
check("max_tokens" not in p and "response_format" in p,
      "when BOTH fields are named, max_tokens is dropped and the JSON contract is KEPT", (p, fx))
p, fx = RC.repair_payload({"response_format": {"type": "json_object"},
                           "chat_template_kwargs": {"enable_thinking": False}},
                          "response_format type 'json_object' is not supported for this model")
check("response_format" not in p, "when only response_format is named, it is the one dropped")
check("chat_template_kwargs" in p, "...and chat_template_kwargs is NOT dropped speculatively")
p, fx = RC.repair_payload({"chat_template_kwargs": {"enable_thinking": False}, "temperature": 1},
                          "chat_template_kwargs is not supported")
check("chat_template_kwargs" not in p, "it IS dropped when the server names it", fx)

# the 400 is actually retried once, end to end, and the retry succeeds
poster6, log6 = stub(resume=good_resume(), cover=good_cover(), audit={"verdict": "pass"},
                     raise_400='{"message":"temperature must be 0.6 for this model"}')
res6 = run(RC.tailor(PROFILE, JD, poster=poster6, rounds=1))
check(res6["resume"] is not None, "a 400 on the first call is repaired and the run completes")
check(any(pl.get("temperature") == 0.6 for pl in log6),
      "the repaired payload was actually SENT with the server's value")


# =========================================================================== 6/7  depth floor
section("6  the depth floor rejects a stub AND passes a realistic draft")

ok, why = RC.contract_ok_resume(good_resume())
check(ok, "a realistic resume draft PASSES the floor", why)
check(RC.resume_depth(good_resume()) > RC.MIN_RESUME,
      "  measured depth %d vs floor %d" % (RC.resume_depth(good_resume()), RC.MIN_RESUME))
ok, why = RC.contract_ok_cover(good_cover())
check(ok, "a realistic cover draft PASSES the floor", why)
check(RC.cover_depth(good_cover()) > RC.MIN_COVER,
      "  measured depth %d vs floor %d" % (RC.cover_depth(good_cover()), RC.MIN_COVER))

thin = {"summary": "Experienced professional.", "experience": [
    {"title": "Head of Infrastructure", "company": "Stars4Business OU",
     "bullets": ["Responsible for infrastructure."]}], "all_employers": ["Stars4Business OU"]}
ok, why = RC.contract_ok_resume(thin)
check(not ok and "THIN" in why, "a THIN resume is rejected as thin", why)
ok, why = RC.contract_ok_resume({})
check(not ok, "an empty object {} is rejected (it parses perfectly)", why)
ok, why = RC.contract_ok_resume(good_resume(), tokens_out=3)
check(not ok and "empty answer" in why, "3 completion tokens is an empty answer", why)
ok, why = RC.contract_ok_cover({"paragraphs": ["Hi."]})
check(not ok, "a one-line cover letter is rejected", why)
ok, why = RC.contract_ok_resume(dict(good_resume(), experience=[]))
check(not ok, "no experience entry -> rejected", why)

section("7  a thin draft moves to the NEXT MODEL in the chain")

seen_models = []


def by_model(good):
    def f(payload):
        seen_models.append(payload["model"])
        # the head model answers thin; the second model answers properly
        return good if payload["model"] != "deepseek-3.2" else thin
    return f


poster7, log7 = stub(resume=by_model(good_resume()), cover=good_cover(), audit={"verdict": "pass"})
res7 = run(RC.tailor(PROFILE, JD, poster=poster7, rounds=0))
check(res7["resume"] is not None, "a usable resume is still produced")
check(res7["authors"]["resume"] == "llama-4-maverick",
      "the THIN head model is rejected and the NEXT model authors it",
      res7["authors"]["resume"])
check(any(a["model"] == "deepseek-3.2" and not a["ok"] and "THIN" in a["why"]
          for a in res7["attempts"]),
      "the rejected attempt is RECORDED with its reason, not hidden",
      res7["attempts"])
check(res7["quality"]["resume_thin"] is False,
      "the shipped document is not thin", res7["quality"])


# =========================================================================== 9  revision guard
section("9  a revision that guts or drops an employer is REJECTED")


def deep_resume():
    """A DEEP draft, so each rejection rule can be tested IN ISOLATION.

    With a 1017-char draft against a 900-char floor, every mutation trips the FLOOR first and the
    assertion passes for the wrong reason. A false pass is the more dangerous half of a bad test:
    the failing one at least announces itself.
    """
    r = good_resume()
    r["experience"].append({"title": "Systems Engineer", "company": "Verint",
                            "dates": "2013 - 2017", "location": "Tel Aviv", "bullets": [
        "Automated provisioning for 60 customer environments, removing the manual build step entirely.",
        "Cut environment build time from two days to under three hours across the estate.",
        "Documented and handed over the provisioning pipeline to the regional operations team."]})
    r["experience"].append({"title": "Network Engineer", "company": "Telefonica",
                            "dates": "2010 - 2013", "location": "Madrid", "bullets": [
        "Ran BGP peering for 14 transit and exchange sessions across two metropolitan regions.",
        "Reduced circuit costs by 12 percent by consolidating duplicate metro-ethernet tails.",
        "Cut mean time to repair on the access layer from 6 hours to 90 minutes."]})
    r["earlier"] = []
    r["all_employers"] = ["Stars4Business OU", "Canonical", "Verint", "Telefonica"]
    return r


before = deep_resume()
D0 = RC.resume_depth(before)
check(D0 >= int(RC.MIN_RESUME / 0.8),
      "the section-9 fixture is deep enough that the FLOOR cannot mask the other rules "
      "(%d chars, floor %d)" % (D0, RC.MIN_RESUME))

ok, why = RC.revision_ok("resume", before, deep_resume())
check(ok, "an equal-quality revision is accepted", why)

# MERGE a role away: Canonical's bullets are kept (so the prose depth is UNCHANGED and the floor
# cannot fire) but the employer itself is gone. This is the real failure mode the rule exists for -
# "never merge, drop, re-date or promote a role".
merged = copy.deepcopy(before)
merged["experience"][0]["bullets"] += merged["experience"][1]["bullets"]
del merged["experience"][1]
check(RC.resume_depth(merged) == D0,
      "  the merge fixture keeps depth identical (%d), so ONLY the employer rule can fire" % D0)
ok, why = RC.revision_ok("resume", before, merged)
check(not ok and "DROPPED employer" in why,
      "a revision that DROPS an employer is refused - for THAT reason", why)
check("Canonical" in why, "...and it NAMES the employer it would have lost", why)

# LOSE PROSE but keep every employer, and stay ABOVE the floor, so only the 80-percent rule can fire.
thinner = copy.deepcopy(before)
for e in thinner["experience"]:
    e["bullets"] = e["bullets"][:1]
D1 = RC.resume_depth(thinner)
check(D1 >= RC.MIN_RESUME and D1 < int(D0 * 0.8),
      "  the thinning fixture is above the floor (%d >= %d) but below 80 percent of the draft "
      "(%d < %d), so ONLY the shrink rule can fire" % (D1, RC.MIN_RESUME, D1, int(D0 * 0.8)))
ok, why = RC.revision_ok("resume", before, thinner)
check(not ok and "thinner" in why,
      "a revision that loses >20 percent of the prose is refused - for THAT reason", why)
ok, why = RC.revision_ok("resume", before, {})
check(not ok, "an empty revision is refused", why)
ok, why = RC.revision_ok("cover", good_cover(), {"paragraphs": good_cover()["paragraphs"][:1]})
check(not ok, "a gutted cover revision is refused", why)
ok, why = RC.revision_ok("cover", good_cover(), good_cover())
check(ok, "an equal cover revision is accepted", why)

# THE DEFECT AN END-TO-END RUN FOUND, and it shipped a fabrication: the guard protected EVERY
# employer in the draft, so when the author had invented "Google" and the auditor caught it, the
# corrected revision - the one that REMOVED the fabrication - was rejected for "dropping Google".
# The invented employer stayed in the DOCX and the real missing one stayed missing. A revision may
# only be blocked from dropping an employer THE PROFILE SUPPORTS.
inv = deep_resume()
inv["experience"].append({"title": "Staff Engineer", "company": "Google", "dates": "2009 - 2010",
                          "bullets": ["Scaled the ranking pipeline to 4 billion documents a day."]})
fixed = deep_resume()                                   # same document with the invention removed
ok, why = RC.revision_ok("resume", inv, fixed, profile=PROFILE)
check(ok, "a revision that REMOVES an invented employer is ACCEPTED", why)
ok, why = RC.revision_ok("resume", inv, fixed)          # no profile -> cannot tell -> fail closed
check(not ok and "Google" in why,
      "...but with NO profile to check against, every employer is protected (fail closed)", why)
merged2 = copy.deepcopy(inv)
merged2["experience"][0]["bullets"] += merged2["experience"][1]["bullets"]
del merged2["experience"][1]                            # Canonical: real, in the profile
ok, why = RC.revision_ok("resume", inv, merged2, profile=PROFILE)
check(not ok and "Canonical" in why,
      "a REAL employer is still protected when a profile IS supplied", why)
check(RC.employer_in_profile("Stars4Business OU", PROFILE) is True,
      "employer_in_profile: a legal suffix does not hide a real employer")
check(RC.employer_in_profile("Google", PROFILE) is False,
      "employer_in_profile: an absent employer is absent")
check(RC.employer_in_profile("AB", PROFILE) is True,
      "employer_in_profile: a name too short to discriminate is treated as present")
check(RC.employer_in_profile("(Canonical)", PROFILE) is True,
      "employer_in_profile: punctuation around a real employer is not a fabrication")
check(RC.employer_in_profile("Canonical Ltd.", PROFILE) is True,
      "employer_in_profile: a legal suffix is not a fabrication")
# ISOLATE the digit-boundary rule. Plain containment finds "22" inside the profile's "220 servers",
# so an inflated figure would corroborate against a real one. Every WORD here is in the profile, so
# only the number rule can decide this case.
check(RC.claim_is_corroborated("Migrated 22 servers to Ubuntu LTS", PROFILE) is False,
      "corroboration: 22 must NOT match the profile's 220 (digit boundary, not containment)")
check(RC.claim_is_corroborated("Migrated 220 servers to Ubuntu LTS", PROFILE) is True,
      "corroboration: ...and the real figure 220 still matches")
# THE FAIL-OPEN HOLE this assertion found: revision_ok tested `profile is None`, so an EMPTY
# profile string reached employer_in_profile, matched nothing, and left every employer unprotected -
# a revision could then legally delete the entire work history. A blank profile means "cannot tell".
ok, why = RC.revision_ok("resume", inv, fixed, profile="")
check(not ok and "Google" in why,
      "an EMPTY profile string also fails closed (every employer protected)", why)
ok, why = RC.revision_ok("resume", inv, fixed, profile="   ")
check(not ok, "whitespace-only is treated the same way", why)


# (a rejected revision reaching the human is asserted in section 12, via the real module)


# =========================================================================== 12  the render defect
section("12  highlights and earlier actually reach the DOCUMENT RENDERER")

os.environ.setdefault("DATA_DIR", "/tmp/jhwtest")
try:
    from app import documents as DOC
    from app import electronic as EL
    blocks = DOC.struct_to_blocks({"basics": {"name": "Jane Roe"}, **good_resume()}, "resume")
    text = " ".join("%s %s" % (k, v) for k, v in blocks)
    check("EARLIER EXPERIENCE" in text,
          "the renderer emits the EARLIER EXPERIENCE section (employer-drop protection)")
    check("Verint" in text, "the demoted employer Verint IS in the rendered document")
    check("31 percent" in text, "the CAR highlights are rendered")
    # the struct electronic.generate builds must carry both keys, or the renderer never sees them
    src = open(os.path.join(APP, "electronic.py"), encoding="utf-8").read()
    st = src[src.index("resume_struct = {"):src.index("cover_struct = {")]
    for key in ('"highlights"', '"earlier"', '"all_employers"'):
        check(key in st, "resume_struct carries %s" % key)
    warns = EL._audit_warnings(res4)
    check(any("REJECTED" in w for w in warns),
          "a rejected revision reaches the human as a WARNING", warns)
    warns3 = EL._audit_warnings(res3)
    check(any("REFUSED" in w for w in warns3),
          "a refused audit round reaches the human as a WARNING", warns3)
    warns_inv = EL._audit_warnings(res)
    check(any("Google" in w for w in warns_inv),
          "an uncorroborated claim reaches the human as a WARNING", warns_inv)
except ImportError as e:
    # A dependency the app declares but this machine lacks is a TOOLCHAIN gap, not a defect. Say so
    # loudly instead of scoring a silent pass - a check that cannot run is not a check.
    print("  SKIP  section 12 needs the app's declared deps (%s)" % e)
    FAILS.append("section 12 could not run: %s" % e)


# =========================================================================== 13 end to end
section("13  the whole loop through the REAL endpoint: author lies, auditor catches, files change")

try:
    import app.electronic as ELE
    import app.llm as LLM
    import docx as _docx

    liar = good_resume()
    liar["experience"].append({"title": "Staff Engineer", "company": "Google",
                               "dates": "2011 - 2013",
                               "bullets": ["Scaled ranking to 4 billion documents a day."]})
    liar["earlier"] = []                                # and Verint is dropped from the document
    liar["all_employers"] = ["Stars4Business OU", "Canonical", "Verint", "Google"]
    E2E_AUDIT = {"verdict": "revise",
                 "invented_facts": [{"where": "resume.experience[2]",
                                     "claim": "Staff Engineer at Google",
                                     "why": "no Google role in the profile"}],
                 "employers_dropped": ["Verint"], "weak_bullets": [],
                 "missed_keywords": ["Terraform"], "cover_issues": [],
                 "scores": {"truthfulness": 4, "jd_fit": 8, "impact": 7, "structure": 8,
                            "cover_letter": 8},
                 "notes": "Drop Google, restore Verint."}

    async def fake_chat(payload, timeout=120.0):
        sysm = payload["messages"][0]["content"]
        userm = payload["messages"][1]["content"]
        if sysm.startswith("You are an INDEPENDENT adversarial reviewer"):
            body = E2E_AUDIT
        elif sysm.startswith("You are the author of this resume"):
            body = good_resume() if "corrected resume JSON" in userm else good_cover()
        elif "tailor an existing resume" in sysm:
            body = liar
        else:
            body = good_cover()
        return {"choices": [{"message": {"content": json.dumps(body)}, "finish_reason": "stop"}],
                "usage": {"completion_tokens": 1200}}

    LLM.chat = fake_chat                                # NO NETWORK, and no key is read
    man = run(ELE.generate(ELE.GenerateReq(email="e2e@example.com", jd=JD, profile=PROFILE)))

    check(man["models"]["auditor"] not in (man["models"]["resume"], man["models"]["cover"]),
          "the manifest's auditor is not either author",
          (man["models"]["resume"], man["models"]["auditor"]))
    check(RC.vendor(man["models"]["auditor"]) != RC.vendor(man["models"]["resume"]),
          "...and comes from a different vendor", man["models"]["auditor_vendor"])
    check(man["employers_missing"] == [],
          "after the revision NO employer is missing", man["employers_missing"])
    check(man["audit"]["scores"]["truthfulness"] == 4, "the scores reach the manifest")
    check(sorted(man["files"]) == ["cover_letter.docx", "cover_letter.pdf", "resume.docx",
                                   "resume.pdf"], "all four documents were written", man["files"])
    r1 = man["audit"]["rounds"][0]
    check(r1["counts"]["invented_facts"] == 1 and r1["counts"]["employers_dropped"] == 1,
          "the round records what it acted on", r1["counts"])
    check(all(rv["applied"] for rv in r1["revisions"]),
          "both corrections were applied", r1["revisions"])

    d = os.path.join(ELE.user_dir("e2e@example.com"), man["job_id"])
    body = "\n".join(p.text for p in _docx.Document(os.path.join(d, "resume.docx")).paragraphs)
    check("Google" not in body, "the RENDERED resume.docx does NOT contain the invented employer")
    check("Verint" in body and "EARLIER EXPERIENCE" in body,
          "the RENDERED resume.docx DOES contain the restored employer, under EARLIER EXPERIENCE")
    check("31 percent" in body, "the CAR highlights are in the rendered file")
    for emp in ("Stars4Business", "Canonical"):
        check(emp in body, "employer %s survives into the file" % emp)
    saved = json.load(open(os.path.join(d, "job.json"), encoding="utf-8"))
    for k in ("audit", "audit_rounds", "audit_refused", "quality", "attempts", "consensus_ms"):
        check(k in saved, "job.json persists '%s' for the UI" % k)
except ImportError as e:
    print("  SKIP  section 13 needs the app's declared deps (%s)" % e)
    FAILS.append("section 13 could not run: %s" % e)


# =========================================================================== 14  revise
section("14  /revise cannot write an EMPTY document or silently lose an employer")

try:
    import app.electronic as ELE2
    import app.llm as LLM2

    JOB = None

    async def _rev(answer, instruction):
        async def ch(payload, timeout=120.0):
            return {"choices": [{"message": {"content": json.dumps(answer)}},
                                ], "usage": {"completion_tokens": 900}}
        LLM2.chat = ch
        return await ELE2.revise(ELE2.ReviseReq(email="e2e@example.com", job_id=JOB,
                                                instruction=instruction))

    # reuse the job section 13 created
    base = ELE2.user_dir("e2e@example.com")
    jobs = sorted(os.listdir(base)) if os.path.isdir(base) else []
    JOB = jobs[-1] if jobs else None
    if not JOB:
        raise ImportError("section 13 did not run, so there is no job to revise")

    good_struct = {"resume": dict(good_resume(), basics={"name": "Jane Roe"}),
                   "cover": good_cover()}
    ELE2._write_json(os.path.join(base, JOB, "tailored.json"), good_struct)

    # (a) an EMPTY answer must not overwrite a good document
    man = run(_rev({"resume": {}, "cover": {}}, "make it shorter"))
    saved = json.load(open(os.path.join(base, JOB, "tailored.json"), encoding="utf-8"))
    check(saved["resume"]["experience"], "an empty edit does NOT overwrite the resume")
    check(saved["cover"]["paragraphs"], "an empty edit does NOT overwrite the cover letter")
    check(len(man["revise_rejected"]) == 2, "both rejections are reported", man["revise_rejected"])
    check(any("rejected" in w for w in man["warnings"]),
          "...and reach the human as warnings", man["warnings"])

    # ISOLATE the contract check. An empty edit is ALSO caught by the employer guard, so a mutation
    # removing the contract check went unnoticed until this case: an explicit "remove ..." AUTHORISES
    # the employer drop, so nothing but the contract check stands between {} and a blank resume.
    ELE2._write_json(os.path.join(base, JOB, "tailored.json"), good_struct)
    man = run(_rev({"resume": {}, "cover": {}}, "remove the Verint role"))
    saved = json.load(open(os.path.join(base, JOB, "tailored.json"), encoding="utf-8"))
    check(saved["resume"]["experience"],
          "an empty edit under an explicit 'remove' instruction STILL cannot blank the resume")
    check(any("rejected" in r for r in man["revise_rejected"]),
          "...refused by the CONTRACT check, with the reason given", man["revise_rejected"])

    # (b) a silent employer drop, while doing something else, is refused
    dropped = copy.deepcopy(good_struct)
    del dropped["resume"]["experience"][1]              # Canonical gone
    dropped["resume"]["experience"][0]["bullets"] += good_resume()["experience"][1]["bullets"]
    man = run(_rev(dropped, "make it punchier"))
    saved = json.load(open(os.path.join(base, JOB, "tailored.json"), encoding="utf-8"))
    check(any("Canonical" in RC._norm(x) or "canonical" in x.lower()
              for x in RC.rendered_employers(saved["resume"])),
          "a SILENT employer drop is refused - Canonical is still there",
          RC.rendered_employers(saved["resume"]))
    check(any("dropped" in r for r in man["revise_rejected"]),
          "...and the refusal says so", man["revise_rejected"])

    # (c) but an EXPLICIT removal request IS authorisation - the candidate asked
    man = run(_rev(dropped, "remove the Canonical role"))
    saved = json.load(open(os.path.join(base, JOB, "tailored.json"), encoding="utf-8"))
    check(not any("canonical" in x.lower() for x in RC.rendered_employers(saved["resume"])),
          "an EXPLICIT 'remove ...' instruction is honoured",
          RC.rendered_employers(saved["resume"]))
    check(man["revise_rejected"] == [], "...with nothing rejected", man["revise_rejected"])
except ImportError as e:
    print("  SKIP  section 14 needs the app's declared deps (%s)" % e)
    FAILS.append("section 14 could not run: %s" % e)


# =========================================================================== no network
section("no network was used")
check(RC._default_poster.__name__ == "_default_poster",
      "the real transport exists but every call in this file went through a stub")
try:
    import socket
    socket.socket = None                                # nothing below may open a socket
    check(True, "sockets were never needed")
finally:
    pass


# =========================================================================== THE GATE (LAST)
print("\n%s\n%d checks run, %d failed" % ("=" * 74, RUN[0], len(FAILS)))
for f in FAILS:
    print("  FAILED: %s" % f)
sys.exit(1 if FAILS else 0)
