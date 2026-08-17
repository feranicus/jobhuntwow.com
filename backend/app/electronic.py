"""electronic.py — the CORE product loop: job description IN -> tailored resume + cover OUT.

Endpoints (mounted by main.py: `app.include_router(electronic.router)`):

    POST /api/electronic/jd                    {text?, url?}            -> parsed JD | needs_paste
    POST /api/electronic/generate              {email, jd, profile}     -> tailored docs + manifest
    GET  /api/electronic/jobs                  ?email=...               -> this user's jobs
    GET  /api/electronic/artifacts/{job_id}    ?email=...               -> list the generated files
    GET  /api/electronic/artifacts/{job_id}/{filename}?email=...        -> download one file

TENANCY: every endpoint takes `email` and resolves to DATA_DIR/users/<sha256(email)[:16]>/<job_id>/.
This module does NOT implement auth (another module owns sessions) - it only scopes storage by the
email it is handed. main.py is expected to inject the authenticated email.

TRUTH RULES (these are the product, not decoration):
  * The LLM may only re-angle, reorder and rephrase facts that are IN the supplied profile.
    It may never add an employer, a date, a degree, a certification, a metric or a tool.
  * A deterministic post-check (`_truth_check`) compares every organisation in the generated
    resume against the profile text and reports anything it cannot find. We surface the
    warning; we never silently "fix" it by rewriting prose.
  * If the profile has no number for a claim, the bullet stays qualitative. No invented metrics.
"""
from __future__ import annotations

import hashlib
import httpx
import io
import json
import os
import re
import sys
import time
import uuid
from typing import Any, Optional

from fastapi import APIRouter, File, HTTPException, Query, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel

from . import documents, jd_ingest, resume_consensus as RC
from .settings import DATA_DIR

router = APIRouter(prefix="/api/electronic", tags=["electronic"])

JOB_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9_-]{0,80}$")
FILE_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,80}$")
ALLOWED_FILES = {"resume.docx", "resume.pdf", "cover_letter.docx", "cover_letter.pdf",
                 "job.json", "tailored.json"}
MIME = {".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".pdf": "application/pdf", ".json": "application/json"}


# --------------------------------------------------------------------------- storage
def user_dir(email: str) -> str:
    """DATA_DIR/users/<sha256(email)[:16]> - stable, non-reversible, filesystem-safe."""
    e = (email or "").strip().lower()
    if not e:
        raise HTTPException(status_code=400, detail="email is required (tenant scope)")
    h = hashlib.sha256(e.encode("utf-8")).hexdigest()[:16]
    return os.path.join(str(DATA_DIR), "users", h)


def job_dir(email: str, job_id: str) -> str:
    if not JOB_ID_RE.match(job_id or ""):
        raise HTTPException(status_code=400, detail="invalid job_id")
    return os.path.join(user_dir(email), job_id)


def _new_job_id(company: str, title: str) -> str:
    slug = re.sub(r"[^a-z0-9]+", "-", ("%s-%s" % (company or "", title or "")).lower()).strip("-")
    slug = (slug[:40] or "job").strip("-")
    return "%s-%s-%s" % (time.strftime("%Y%m%d-%H%M%S"), slug, uuid.uuid4().hex[:6])


def _write_json(path: str, data: dict) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


# --------------------------------------------------------------------------- LLM plumbing
# ONE HOME per rule. `_json`, `_profile_text`, `_missing_employers` and `_truth_check` are now
# implemented in resume_consensus (which needs them for its own deterministic guardrails) and
# aliased here. Two copies of a truth rule is precisely how the two drift apart - the sibling
# project paid for that four times over with one model chain living in four places.
_json = RC.json_loose


async def _ask(system: str, user: str, max_tokens: int) -> dict:
    """One JSON answer, walking the MEASURED model chain. Returns {'_error': ...} rather than raising.

    It used to resolve the roles "content" -> "content_fb", i.e. `deepseek-v4-pro` with
    `deepseek-3.2` behind it. `deepseek-v4-pro` appears in NO measurement in this repository:
    model-selection-toolkit/MODEL_SELECTION.md measured `deepseek-3.2` and `llama-4-maverick` as
    the winners, and neither names a v4-pro. An unverified id is exactly the phantom-model class the
    sibling project already paid for ("DeepSeek V4 Flash" was a marketing name that returned 404 on
    every call and silently degraded every run). So this path now uses the same measured chain the
    consensus author uses - four vendors, all of them ids we can point at evidence for.
    NOT VERIFIED HERE: whether `deepseek-v4-pro` exists on the account. That needs the live catalog
    (`/api/models`). The `content` role is left untouched in llm.py because the agent uses it too.
    """
    last = None
    for m in RC.chain():
        try:
            txt, _usage, _fin = await RC.call_model(m, system, user, max_tokens=max_tokens,
                                                    timeout=RC.TIMEOUT)
            d = RC.json_loose(txt)
            if d:
                d["_model_role"] = m
                return d
            last = "empty/unparseable answer from %s" % m
        except Exception as e:
            last = "%s: %s" % (type(e).__name__, e)
    return {"_error": last or "no model answered"}


# --------------------------------------------------------------------------- prompts
# The prompts, the recruiter rules and the truth rules now live in resume_consensus, because the
# AUTHOR, the AUDITOR and the REVISION step all have to read the same text - and a prompt is a
# string that reaches a human (via the model), so it may not exist in two versions. These aliases
# keep the names documented in docs/TAILOR_LOGIC.md resolvable.
_TRUTH = RC._TRUTH
_RESUME_SYS = RC.RESUME_SYS
_COVER_SYS = RC.COVER_SYS
RESUME_RULES = RC.RESUME_RULES
COVER_RULES = RC.COVER_RULES
_resume_user = RC.resume_user
_cover_user = RC.cover_user


# --------------------------------------------------------------------------- profile handling
_profile_text = RC.profile_text


def _pget(profile: Any, key: str, default):
    return profile.get(key, default) if isinstance(profile, dict) else default


_RE_EMAIL = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_RE_PHONE = re.compile(r"(?:\+\d{1,3}[\s./-]?)?(?:\(?\d{2,4}\)?[\s./-]?){2,5}\d{2,4}")
_RE_LI    = re.compile(r"(?:https?://)?(?:[\w-]+\.)?linkedin\.com/in/[\w%-]+/?", re.I)
_RE_SITE  = re.compile(r"(?:https?://)?(?:www\.)?[\w-]+\.[a-z]{2,}(?:/[\w%./-]*)?", re.I)
_RE_LABEL = re.compile(r"^\s*-?\s*([a-z_]+)\s*:\s*(.+?)\s*$", re.I)


def _contacts_from_text(text: str) -> dict:
    """Pull the contact block out of a REAL resume (a PDF/DOCX has no 'email:' labels).

    Without this the header rendered as "Candidate" with an empty contact line — the documents
    were unusable no matter how good the prose was.
    """
    out, head = {}, [l.strip() for l in (text or "").splitlines() if l.strip()][:12]
    blob = "\n".join(head)

    m = _RE_EMAIL.search(blob)
    if m:
        out["email"] = m.group(0)
    m = _RE_LI.search(blob)
    if m:
        out["linkedin"] = m.group(0).rstrip("/")
    for line in head:
        cand = line
        if out.get("email"):
            cand = cand.replace(out["email"], " ")
        if out.get("linkedin"):
            cand = cand.replace(out["linkedin"], " ")
        m = _RE_PHONE.search(cand)
        # a phone needs enough digits; postcodes and years must not match
        if m and len(re.sub(r"\D", "", m.group(0))) >= 9:
            out.setdefault("phone", m.group(0).strip(" |,;"))
    for line in head:
        for w in _RE_SITE.findall(line):
            lw = w.lower()
            if ("linkedin.com" in lw or "@" in w or lw.endswith((".jpg", ".png"))
                    or (out.get("email") and w in out["email"])):
                continue
            if "." in w and " " not in w and len(w) > 4:
                out.setdefault("website", w.rstrip("/.,;|"))
                break
        if out.get("website"):
            break

    # Name = the first line that is not a contact line and looks like a person's name.
    for line in head:
        if any(t in line for t in ("@", "http", "linkedin.com")) or _RE_PHONE.search(line):
            continue
        words = line.replace(",", " ").split()
        if 1 < len(words) <= 5 and sum(w[:1].isupper() for w in words) >= 2 and len(line) < 60:
            out.setdefault("name", line.strip())
            break
    # Headline = the line right after the name, if it reads like a title.
    if out.get("name") and out["name"] in head:
        i = head.index(out["name"])
        if i + 1 < len(head):
            nxt = head[i + 1]
            if "@" not in nxt and not _RE_PHONE.search(nxt) and 3 < len(nxt) < 140:
                out.setdefault("headline", nxt.strip())
    # Location = the contact line minus phone/email/links (usually "City, Street, ZIP, Country").
    for line in head:
        if _RE_EMAIL.search(line) or _RE_PHONE.search(line):
            part = line
            for k in ("email", "phone", "linkedin", "website"):
                if out.get(k):
                    part = part.replace(out[k], " ")
            part = " ".join(p.strip(" |,;·•") for p in part.split("|"))
            part = re.sub(r"\s{2,}", " ", part).strip(" |,;·•-")
            if len(part) > 6 and any(c.isalpha() for c in part):
                out.setdefault("location", part)
            break
    return out


def _basics(profile: Any) -> dict:
    """Pull the contact block out of a dict profile, or out of markdown 'key: value' lines."""
    keys = ("name", "full_name", "headline", "title", "email", "phone", "location", "city",
            "country", "linkedin", "website", "portfolio")
    out = {}
    if isinstance(profile, dict):
        src = profile.get("basics") if isinstance(profile.get("basics"), dict) else profile
        for k in keys:
            v = src.get(k)
            if isinstance(v, (str, int, float)) and str(v).strip():
                out[k] = str(v).strip()
    else:
        txt = str(profile or "")
        for line in txt.split("\n"):
            m = _RE_LABEL.match(line)
            if m and m.group(1).lower() in keys:
                out.setdefault(m.group(1).lower(), m.group(2).strip().strip('"'))
        # A real resume has no "email:" labels -> parse the contact block out of the text itself.
        for k, v in _contacts_from_text(txt).items():
            out.setdefault(k, v)
    if "name" not in out and "full_name" in out:
        out["name"] = out["full_name"]
    if "headline" not in out and "title" in out:
        out["headline"] = out["title"]
    if "location" not in out:
        loc = " ".join(x for x in (out.get("city"), out.get("country")) if x)
        if loc:
            out["location"] = loc
    return out


_missing_employers = RC.missing_employers
_truth_check = RC.truth_check


def _audit_warnings(con: dict) -> list:
    """Turn the consensus record into warnings a human reads. Surface, never bury.

    An adversarial reviewer whose findings are stored in a JSON blob nobody renders is decoration.
    These are the four things the candidate must know BEFORE the document goes out under their name:
    a claim the reviewer could not tie to the profile, a revision we refused to accept, an audit
    round we refused entirely, and a review that never happened at all.
    """
    out = []
    aud = con.get("audit") or {}
    if aud.get("skipped"):
        out.append("No independent review was performed: %s" % aud["skipped"])
    for rnd in (aud.get("rounds") or []):
        if rnd.get("error"):
            out.append("Independent review round %s did not complete (%s) - the draft is "
                       "unreviewed." % (rnd.get("round"), rnd["error"]))
        if rnd.get("refused_all"):
            out.append("The independent review was REFUSED and nothing was changed: %s. Read the "
                       "draft yourself before sending." % rnd.get("refused_reason", ""))
        for f in (rnd.get("kept") or {}).get("invented_facts") or []:
            out.append("Independent review could not find this in your profile: \"%s\" (%s) - "
                       "verify or remove it." % (str(f.get("claim", ""))[:180],
                                                 str(f.get("where", ""))[:60]))
        for r in (rnd.get("revisions") or []):
            if not r.get("applied"):
                out.append("A correction to the %s was REJECTED and the earlier draft kept: %s"
                           % (r.get("doc"), r.get("why", "")))
    q = con.get("quality") or {}
    if q.get("resume_thin"):
        out.append("The resume is thin (%s chars of prose, floor %s) - every model in the chain "
                   "under-delivered." % (q.get("resume_chars"), q.get("resume_floor")))
    if q.get("cover_thin"):
        out.append("The cover letter is thin (%s chars, floor %s)."
                   % (q.get("cover_chars"), q.get("cover_floor")))
    seen, uniq = set(), []
    for w in out:
        if w not in seen:
            seen.add(w)
            uniq.append(w)
    return uniq


# --------------------------------------------------------------------------- API models
class JDReq(BaseModel):
    text: Optional[str] = ""
    url: Optional[str] = ""


class GenerateReq(BaseModel):
    email: str
    jd: Any = None                 # the dict from /jd, or a raw JD string
    profile: Any = None            # dict or markdown text (candidate.md)
    job_id: Optional[str] = ""
    answers: Any = None            # {gap question: candidate's answer} from the gap Q&A
    evidence: Any = None           # [{name, text}] extracted from uploaded portfolio/articles
    links: Any = None              # ["https://..."] portfolio / publication URLs
    use_photo: bool = False


# --------------------------------------------------------------------------- endpoints
def extract_text(name: str, blob: bytes) -> tuple[str, str]:
    """Return (text, kind) from an uploaded resume/profile.

    The browser cannot read a PDF or DOCX - FileReader hands back raw bytes, which is how a
    resume arrived as "%PDF-1.4 ..." and the truth-check reported "Profile is
    encrypted/unreadable". Extraction has to happen HERE.
    """
    low = (name or "").lower()
    if low.endswith(".pdf"):
        from pypdf import PdfReader
        rd = PdfReader(io.BytesIO(blob))
        if getattr(rd, "is_encrypted", False):
            try:
                rd.decrypt("")           # many "encrypted" PDFs use an empty owner password
            except Exception:
                raise HTTPException(400, "That PDF is password-protected - save an unlocked copy.")
        pages = [(p.extract_text() or "") for p in rd.pages]
        return "\n".join(pages).strip(), "pdf"
    if low.endswith(".docx"):
        import docx
        d = docx.Document(io.BytesIO(blob))
        parts = [p.text for p in d.paragraphs]
        for t in d.tables:                # resumes very often put dates/roles in tables
            for row in t.rows:
                parts.append("\t".join(c.text for c in row.cells))
        return "\n".join(x for x in parts if x.strip()).strip(), "docx"
    if low.endswith(".doc"):
        raise HTTPException(400, "Legacy .doc is not supported - save as .docx or PDF.")
    return blob.decode("utf-8", "replace").strip(), "text"


@router.post("/profile/upload")
async def profile_upload(file: UploadFile = File(...)):
    """PDF / DOCX / TXT / MD in -> plain text out, so the tailor never sees raw bytes."""
    blob = await file.read()
    if len(blob) > 8 * 1024 * 1024:
        raise HTTPException(400, "File is larger than 8 MB.")
    text, kind = extract_text(file.filename or "", blob)
    words = len(text.split())
    # Floor exists so we never tailor from garbage (a raw-bytes PDF reads as ~0 real words), but
    # it must not reject a genuinely short profile -- and the message must match the actual cause.
    if words < 25:
        if kind == "pdf":
            hint = ("This PDF appears to be a scan/image: there is no selectable text in it. "
                    "Export a text PDF, or paste your resume below.")
        else:
            hint = "Paste more detail below - there is not enough here to tailor honestly."
        raise HTTPException(400, f"Only {words} words could be read from {file.filename!r}. {hint}")
    return {"text": text, "kind": kind, "chars": len(text), "words": words,
            "filename": file.filename}


@router.post("/evidence/upload")
async def evidence_upload(email: str = Query(...), file: UploadFile = File(...)):
    """A supporting artifact: project portfolio, case study, published article, reference letter.

    Same extraction path as the profile (PDF/DOCX/TXT/MD). The text is returned to the browser and
    sent back with the next generate() call, so it becomes evidence the model may cite - it is NOT
    silently merged into the stored profile.
    """
    blob = await file.read()
    if len(blob) > 8 * 1024 * 1024:
        raise HTTPException(400, "File is larger than 8 MB.")
    text, kind = extract_text(file.filename or "", blob)
    if len(text.split()) < 10:
        raise HTTPException(400, f"Almost no text could be read from {file.filename!r}.")
    return {"name": file.filename, "kind": kind, "words": len(text.split()), "text": text}


async def _fetch_link(url: str) -> dict:
    """Best-effort: read an article/portfolio URL. Many sites block bots - we never pretend."""
    try:
        async with httpx.AsyncClient(timeout=jd_ingest.TIMEOUT, follow_redirects=True) as c:
            r = await c.get(url, headers={"User-Agent": jd_ingest.UA})
        if r.status_code >= 400:
            return {"url": url, "ok": False, "note": f"HTTP {r.status_code}", "text": ""}
        txt = jd_ingest.html_to_text(r.text)[:6000]
        return {"url": url, "ok": bool(txt.strip()), "text": txt,
                "note": "" if txt.strip() else "no readable text"}
    except Exception as e:
        return {"url": url, "ok": False, "note": type(e).__name__, "text": ""}


@router.post("/photo/upload")
async def photo_upload(email: str = Query(...), file: UploadFile = File(...)):
    """Store a CV photo for this user. Optional by design.

    A photo is customary on a German/Austrian/Swiss CV and is a liability in US/UK screening
    (and some ATS parsers choke on images), so the UI must present it as opt-in.
    """
    blob = await file.read()
    if len(blob) > 4 * 1024 * 1024:
        raise HTTPException(400, "Photo is larger than 4 MB.")
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in (".jpg", ".jpeg", ".png"):
        raise HTTPException(400, "Use a JPG or PNG photo.")
    try:
        from PIL import Image
        im = Image.open(io.BytesIO(blob))
        im.verify()                                  # reject anything that is not a real image
    except Exception:
        raise HTTPException(400, "That file is not a readable image.")
    d = user_dir(email)
    os.makedirs(d, exist_ok=True)
    path = os.path.join(d, "photo" + (".png" if ext == ".png" else ".jpg"))
    with open(path, "wb") as f:
        f.write(blob)
    return {"ok": True, "photo": os.path.basename(path), "bytes": len(blob)}


def _photo_path(email: str) -> str:
    for name in ("photo.jpg", "photo.png"):
        p = os.path.join(user_dir(email), name)
        if os.path.isfile(p):
            return p
    return ""


@router.post("/jd")
async def parse_jd(req: JDReq):
    """Job description IN. Pasted text always works; URLs are best-effort by tier."""
    return await jd_ingest.ingest(text=req.text or "", url=req.url or "")


@router.post("/generate")
async def generate(req: GenerateReq):
    """Tailor resume + cover letter to the JD and write DOCX + PDF for both."""
    t0 = time.time()

    jd = req.jd
    if isinstance(jd, str):
        jd = jd_ingest.from_text(jd)
    if not isinstance(jd, dict):
        raise HTTPException(status_code=400, detail="jd must be an object or a string")
    if jd.get("status") in ("needs_paste", "needs_local_fetch"):
        raise HTTPException(status_code=400, detail=jd.get("note") or "JD could not be fetched")
    if not str(jd.get("text") or "").strip():
        raise HTTPException(status_code=400,
                            detail="jd.text is empty - parse it first via POST /api/electronic/jd")

    if not req.profile:
        raise HTTPException(status_code=400,
                            detail="profile is required - we never invent candidate facts")
    ptxt = _profile_text(req.profile)

    # Gap answers are FACTS THE CANDIDATE CONFIRMED (e.g. "I ran projects for Aldi and AON" —
    # true, but absent from the uploaded CV). They are appended to the profile so the model may
    # use them; the truth-check treats the profile+answers as the permitted source, so this does
    # not weaken the no-invention rule.
    answered = []
    if req.answers:
        items = req.answers.items() if isinstance(req.answers, dict) else [
            (a.get("q", ""), a.get("a", "")) for a in req.answers if isinstance(a, dict)]
        for q, a in items:
            if str(a or "").strip():
                answered.append(f"- {str(q).strip()} -> {str(a).strip()}")
    if answered:
        ptxt += ("\n\nADDITIONAL FACTS CONFIRMED BY THE CANDIDATE (treat as true, use where "
                 "relevant to the target job):\n" + "\n".join(answered))

    # Supporting artifacts: portfolio, case studies, published articles, reference letters.
    # Capped per item so one long PDF cannot crowd the JD out of the prompt window.
    ev_used, link_used, link_notes = 0, 0, []
    for e in (req.evidence or []):
        if isinstance(e, dict) and str(e.get("text") or "").strip():
            ptxt += ("\n\nSUPPORTING DOCUMENT PROVIDED BY THE CANDIDATE (%s):\n%s"
                     % (str(e.get("name") or "attachment")[:120], str(e["text"])[:6000]))
            ev_used += 1
    for u in (req.links or []):
        u = str(u or "").strip()
        if not u.startswith(("http://", "https://")):
            continue
        got = await _fetch_link(u)
        if got["ok"]:
            ptxt += "\n\nLINK PROVIDED BY THE CANDIDATE (%s):\n%s" % (u, got["text"])
            link_used += 1
        else:
            # Say so instead of silently ignoring it - the candidate thinks it was read.
            link_notes.append("%s could not be read (%s)" % (u, got["note"]))
    if len(ptxt.strip()) < 60:
        raise HTTPException(status_code=400,
                            detail="profile is too thin to tailor from (we need the candidate's "
                                   "real experience; we will not invent it)")

    jid = req.job_id or _new_job_id(jd.get("company", ""), jd.get("title", ""))
    if not JOB_ID_RE.match(jid):
        raise HTTPException(status_code=400, detail="invalid job_id")
    outdir = job_dir(req.email, jid)
    os.makedirs(outdir, exist_ok=True)

    # 4-MODEL CONSENSUS. One model AUTHORS, a DIFFERENT model from a DIFFERENT VENDOR AUDITS
    # against the JD + profile, then the author revises - bounded rounds. This replaced a single
    # call per document whose answer WAS the document: nothing measured whether the draft was
    # thin, nothing checked it against the JD, and a well-formed nothing rendered as a finished
    # resume. resume_consensus owns the chain, the depth floors and the guardrail that an audit can
    # never empty or gut a document. It never raises: partial success is legible.
    con = await RC.tailor(ptxt, jd)
    tailored = con["resume"] or {}
    cover = con["cover"] or {}

    errors = dict(con["errors"])
    if not con["resume"] and not con["cover"]:
        # EVERY model in the chain failed the contract or the depth floor for BOTH documents.
        # Report what each one actually said instead of a generic 502 - "the model failed" is a
        # conclusion, not an observation.
        raise HTTPException(status_code=502,
                            detail={"error": "no model in the chain produced a usable document",
                                    "per_document": errors, "attempts": con["attempts"]})

    basics = _basics(req.profile)
    # MEASURED DEFECT, FIXED HERE: documents.py renders `highlights` (line 126) and `earlier`
    # (line 157), and this struct never set either. So the 3-5 CAR highlights that are supposed to
    # LEAD the resume - the top third a recruiter reads in the first 7 seconds - were requested
    # from the model, returned, and thrown away. Worse, `earlier` is the ONLY mechanism that keeps
    # a demoted employer in the document, so every older employer was silently absent from the
    # rendered file while `_missing_employers(tailored)` inspected the model's ANSWER and reported
    # clean. The manifest said the employers were there; the DOCX did not have them.
    # `all_employers` is carried so the employer-drop check runs against what was RENDERED.
    resume_struct = {
        "basics": basics,
        "summary": tailored.get("summary") or _pget(req.profile, "summary", ""),
        "highlights": tailored.get("highlights") or _pget(req.profile, "highlights", []),
        "skills": tailored.get("skills") or _pget(req.profile, "skills", []),
        "experience": tailored.get("experience") or _pget(req.profile, "experience", []),
        "earlier": tailored.get("earlier") or [],
        "all_employers": tailored.get("all_employers") or [],
        "education": _pget(req.profile, "education", []),
        "certifications": _pget(req.profile, "certifications", []),
        "languages": _pget(req.profile, "languages", []),
    }
    cover_struct = {
        "basics": basics,
        "company": jd.get("company", ""),
        "job_title": jd.get("title", ""),
        "salutation": cover.get("salutation") or "Hiring Team",
        "paragraphs": cover.get("paragraphs") or [],
        "closing": cover.get("closing") or "Sincerely,",
    }

    photo = _photo_path(req.email) if req.use_photo else ""
    written = documents.write_all(resume_struct, cover_struct, outdir, photo=photo)
    if written.get("errors"):
        errors.update(written["errors"])
    if not written.get("files"):
        raise HTTPException(status_code=500, detail="document rendering failed: %s" % errors)

    manifest = {
        "job_id": jid,
        "email": req.email,          # daily_report attributes the job to a person
        "created": int(t0),
        "elapsed_ms": int((time.time() - t0) * 1000),
        "jd": {k: jd.get(k, "") for k in ("title", "company", "location", "source", "url", "note")},
        # WHO wrote it and WHO reviewed it. A reader must be able to see that the auditor was not
        # the author, without taking our word for it.
        "models": {"resume": con["authors"].get("resume"), "cover": con["authors"].get("cover"),
                   "auditor": con["audit"].get("auditor"),
                   "auditor_vendor": con["audit"].get("auditor_vendor")},
        "attempts": con["attempts"],
        "files": sorted(os.path.basename(p) for p in written["files"].values()),
        "keywords_matched": tailored.get("keywords_matched") or [],
        "gaps": tailored.get("gaps") or [],
        # The employer-drop check now runs on the RENDERED struct, not on the model's answer -
        # the struct is what became the DOCX, and those were not the same thing (see above).
        "warnings": (_truth_check(resume_struct, req.profile)
                     + ["Employer missing from the resume: %s" % m
                        for m in _missing_employers(resume_struct)]
                     + _audit_warnings(con)),
        "employers_missing": _missing_employers(resume_struct),
        # THE CONSENSUS RECORD. Never hidden: the auditor's corroborated findings, what it flagged
        # that we REFUSED to act on, and every rejected revision. Hiding these would defeat the
        # point of having an adversarial reviewer at all.
        "audit": con["audit"],
        "audit_rounds": con["audit_rounds"],
        "audit_refused": con["audit_refused"],
        "quality": con["quality"],
        "consensus_ms": con["elapsed_ms"],
        "answers_used": len(answered),
        "evidence_used": ev_used,
        "links_used": link_used,
        "link_notes": link_notes,
        "photo_used": bool(photo),
        "errors": errors,
    }
    _write_json(os.path.join(outdir, "job.json"), manifest)
    _write_json(os.path.join(outdir, "tailored.json"),
                {"resume": resume_struct, "cover": cover_struct})
    return manifest


class ReviseReq(BaseModel):
    email: str
    job_id: str
    instruction: str
    use_photo: bool = False


@router.post("/revise")
async def revise(req: ReviseReq):
    """Iterate on documents already built: 'add Canonical back', 'shorter', 'more formal'.

    Editing the stored STRUCT (not regenerating from the JD) means the tailoring survives and the
    change is cheap. The truth rules still apply: the model may only rearrange, trim or reword
    what is already there.
    """
    outdir = job_dir(req.email, req.job_id)
    tp = os.path.join(outdir, "tailored.json")
    if not os.path.isfile(tp):
        raise HTTPException(404, "No such job for this account.")
    cur = json.load(open(tp, encoding="utf-8"))
    instruction = (req.instruction or "").strip()
    if not instruction:
        raise HTTPException(400, "Say what you want changed.")

    system = ("You edit an existing resume/cover-letter JSON. Apply ONLY the requested change. "
              "Never invent facts, never delete an employer, keep every key that exists. "
              "Return the SAME JSON shape, nothing else.")
    user = ("CURRENT DOCUMENT JSON:\n%s\n\nREQUESTED CHANGE:\n%s\n\n"
            "Return the complete updated JSON with keys 'resume' and 'cover'."
            % (json.dumps(cur)[:24000], instruction[:2000]))
    got = await _ask(system, user, 6500)
    if got.get("_error"):
        raise HTTPException(502, "The model could not apply that change: %s" % got["_error"])

    # A REVISE COULD PREVIOUSLY WRITE AN EMPTY DOCUMENT. Any dict was accepted, so a model that
    # answered {"resume": {}} overwrote a good resume with nothing and re-rendered a blank DOCX.
    # The same guardrail the consensus uses applies here: the edit must pass the contract, and it
    # may not silently lose an employer.
    prev_resume = cur.get("resume") or {}
    prev_cover = cur.get("cover") or {}
    resume_struct, cover_struct = prev_resume, prev_cover
    rejected = []
    cand_r = got.get("resume")
    if isinstance(cand_r, dict):
        cand_r = RC.normalise_resume(cand_r)
        # basics/education/certifications/languages are NOT the model's to rewrite - they come from
        # the profile, so they are carried over rather than trusted from the edit.
        for k in ("basics", "education", "certifications", "languages"):
            cand_r.setdefault(k, prev_resume.get(k))
        ok, why = RC.contract_ok_resume(cand_r)
        if not ok:
            rejected.append("resume edit rejected: %s - the previous version was kept" % why)
        else:
            lost = [c for c in RC.rendered_employers(prev_resume)
                    if not any(RC._norm(c) in RC._norm(x) or RC._norm(x) in RC._norm(c)
                               for x in RC.rendered_employers(cand_r))]
            # An EXPLICIT instruction to remove something IS authorisation - the candidate asked.
            # A drop that happens while doing something else is the bug this rule exists for.
            asked = bool(re.search(r"\b(remove|drop|delete|take out|kill)\b", instruction, re.I))
            if lost and not asked:
                rejected.append("resume edit rejected: it would have dropped %s and you did not "
                                "ask to remove anything - the previous version was kept"
                                % ", ".join(sorted(set(lost))[:6]))
            else:
                resume_struct = cand_r
    cand_c = got.get("cover")
    if isinstance(cand_c, dict):
        cand_c = RC.normalise_cover(cand_c)
        cand_c.setdefault("basics", prev_cover.get("basics"))
        for k in ("company", "job_title"):
            cand_c.setdefault(k, prev_cover.get(k))
        ok, why = RC.contract_ok_cover(cand_c)
        if not ok:
            rejected.append("cover-letter edit rejected: %s - the previous version was kept" % why)
        else:
            cover_struct = cand_c
    for r in rejected:
        print("[revise] %s" % r, file=sys.stderr)
    photo = _photo_path(req.email) if req.use_photo else ""
    written = documents.write_all(resume_struct, cover_struct, outdir, photo=photo)
    _write_json(tp, {"resume": resume_struct, "cover": cover_struct})

    manifest = {}
    mp = os.path.join(outdir, "job.json")
    if os.path.isfile(mp):
        manifest = json.load(open(mp, encoding="utf-8"))
    manifest.update({
        "files": sorted(os.path.basename(x) for x in written["files"].values()),
        "errors": written["errors"],
        "revised": instruction[:200],
        "revise_model": got.get("_model_role"),
        "revise_rejected": rejected,
        "employers_missing": _missing_employers(resume_struct),
        "warnings": (["Employer missing from the resume: %s" % m
                      for m in _missing_employers(resume_struct)] + rejected),
        "photo_used": bool(photo),
    })
    _write_json(mp, manifest)
    return manifest


class AppliedReq(BaseModel):
    email: str
    job_id: Optional[str] = ""
    url: Optional[str] = ""
    company: Optional[str] = ""
    title: Optional[str] = ""
    status: Optional[str] = "applied"
    note: Optional[str] = ""


@router.post("/applied")
def mark_applied(req: AppliedReq):
    """The local apply engine reports what it actually submitted, so the Pipeline is real.

    Writes into the job's own manifest when we know the job_id (documents were generated here),
    otherwise creates a stub record - an application driven from a pasted URL still belongs in
    the funnel.
    """
    d = user_dir(req.email)
    os.makedirs(d, exist_ok=True)
    jid = req.job_id or _new_job_id(req.company or "", req.title or "")
    outdir = job_dir(req.email, jid)
    os.makedirs(outdir, exist_ok=True)
    mp = os.path.join(outdir, "job.json")
    man = {}
    if os.path.isfile(mp):
        try:
            man = json.load(open(mp, encoding="utf-8"))
        except Exception:
            man = {}
    man.setdefault("job_id", jid)
    man.setdefault("email", req.email)
    man.setdefault("created", int(time.time()))
    jd = man.get("jd") or {}
    jd.update({k: v for k, v in (("company", req.company), ("title", req.title),
                                 ("url", req.url)) if v})
    man["jd"] = jd
    man["applied"] = {"at": int(time.time()), "status": req.status or "applied",
                      "note": (req.note or "")[:500], "url": req.url or ""}
    _write_json(mp, man)
    return {"ok": True, "job_id": jid, "status": man["applied"]["status"]}


@router.get("/jobs")
def list_jobs(email: str = Query(...)):
    """Every job this user has generated, newest first."""
    base = user_dir(email)
    if not os.path.isdir(base):
        return {"count": 0, "jobs": []}
    jobs = []
    for name in sorted(os.listdir(base), reverse=True):
        if not JOB_ID_RE.match(name) or not os.path.isdir(os.path.join(base, name)):
            continue
        item = {"job_id": name}
        mf = os.path.join(base, name, "job.json")
        if os.path.exists(mf):
            try:
                with open(mf, encoding="utf-8") as f:
                    m = json.load(f)
                jd = m.get("jd", {}) or {}
                item.update({"jd": jd, "created": m.get("created"),
                             "title": jd.get("title", ""), "company": jd.get("company", ""),
                             "files": m.get("files", []), "warnings": m.get("warnings", []),
                             # the Pipeline needs to know what was actually SENT, not just built
                             "applied": m.get("applied")})
            except Exception:
                item["note"] = "manifest unreadable"
        jobs.append(item)
    return {"count": len(jobs), "jobs": jobs}


@router.get("/artifacts/{job_id}")
def artifacts(job_id: str, email: str = Query(...)):
    """List the generated files for one job (with sizes + download URLs)."""
    d = job_dir(email, job_id)
    if not os.path.isdir(d):
        raise HTTPException(status_code=404, detail="job not found")
    files = []
    for name in sorted(os.listdir(d)):
        p = os.path.join(d, name)
        if not os.path.isfile(p) or not FILE_RE.match(name):
            continue
        files.append({"name": name, "bytes": os.path.getsize(p),
                      "url": "/api/electronic/artifacts/%s/%s" % (job_id, name)})
    manifest = {}
    mp = os.path.join(d, "job.json")
    if os.path.exists(mp):
        try:
            with open(mp, encoding="utf-8") as f:
                manifest = json.load(f)
        except Exception:
            manifest = {"note": "manifest unreadable"}
    return {"job_id": job_id, "files": files, "manifest": manifest}


@router.get("/artifacts/{job_id}/{filename}")
def artifact(job_id: str, filename: str, email: str = Query(...)):
    """Download one generated file. Path traversal is impossible: the name is allow-listed."""
    if not FILE_RE.match(filename or "") or filename not in ALLOWED_FILES:
        raise HTTPException(status_code=400, detail="unknown artifact")
    d = job_dir(email, job_id)
    p = os.path.join(d, filename)
    # belt and braces: the resolved path must still be inside the user's job dir
    if os.path.commonpath([os.path.realpath(p), os.path.realpath(d)]) != os.path.realpath(d):
        raise HTTPException(status_code=400, detail="unknown artifact")
    if not os.path.isfile(p):
        raise HTTPException(status_code=404, detail="artifact not found")
    ext = os.path.splitext(filename)[1].lower()
    return FileResponse(p, media_type=MIME.get(ext, "application/octet-stream"), filename=filename)
