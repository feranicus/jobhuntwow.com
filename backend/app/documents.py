"""documents.py — render a tailored resume / cover letter to real DOCX and PDF files.

Design
------
Both writers consume ONE intermediate form: a flat list of BLOCKS

    ("h1",      text)   document title (the candidate name)
    ("contact", text)   the contact line under the name
    ("h2",      text)   section heading (SUMMARY / EXPERIENCE / ...)
    ("h3",      text)   job title line  ("Role - Company")
    ("meta",    text)   dates / location line under an h3
    ("p",       text)   paragraph
    ("li",      text)   bullet
    ("spacer",  "")     vertical gap

So DOCX and PDF can never drift: `struct_to_blocks()` / `markdown_to_blocks()` produce the
blocks once, and `_write_docx()` / `_write_pdf()` are dumb renderers.

ATS-safety rules baked in (these are the reason ATS parsers choke on pretty resumes):
  * NO tables, NO text boxes, NO columns, NO headers/footers, NO images.
  * One standard font (Calibri in DOCX / Helvetica in PDF), single column, left aligned.
  * Real Word heading styles + real list bullets so the XML carries structure.
  * Section headings use the conventional ATS keywords (EXPERIENCE, EDUCATION, SKILLS).

Dependencies are pure-Python on purpose: python-docx and reportlab. NOT weasyprint or
wkhtmltopdf - those need OS packages (cairo/pango/Qt) we do not ship in the image.

Nothing here ever raises on a missing optional field: every getter is defensive.
"""
from __future__ import annotations

import datetime as _dt
import os
import re
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer

FONT_DOCX = "Calibri"
FONT_PDF = "Helvetica"
FONT_PDF_B = "Helvetica-Bold"
ACCENT = RGBColor(0x1F, 0x36, 0x4D)   # dark navy - prints near-black, safe for ATS


# --------------------------------------------------------------------------- coercion
def _s(v: Any) -> str:
    """Coerce anything to a clean string. Never raises."""
    if v is None:
        return ""
    if isinstance(v, (list, tuple)):
        return " - ".join(_s(x) for x in v if _s(x))
    if isinstance(v, dict):
        return " - ".join(_s(x) for x in v.values() if _s(x))
    s = str(v)
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    return s.strip()


def _list(v: Any) -> list:
    if v is None:
        return []
    if isinstance(v, (list, tuple)):
        return [x for x in v if x is not None]
    return [v]


def _join(parts, sep=" | ") -> str:
    return sep.join(p for p in (_s(x) for x in parts) if p)


# --------------------------------------------------------------------------- struct -> blocks
def _contact_line(d: dict) -> str:
    b = d.get("basics") if isinstance(d.get("basics"), dict) else d
    return _join([b.get("location"), b.get("phone"), b.get("email"),
                  b.get("linkedin"), b.get("website") or b.get("portfolio")])


def _name_of(d: dict) -> str:
    b = d.get("basics") if isinstance(d.get("basics"), dict) else d
    return _s(b.get("name") or b.get("full_name")) or "Candidate"


def _headline_of(d: dict) -> str:
    b = d.get("basics") if isinstance(d.get("basics"), dict) else d
    return _s(b.get("headline") or b.get("title"))


def _resume_blocks(d: dict) -> list:
    out = [("h1", _name_of(d))]
    hl = _headline_of(d)
    if hl:
        out.append(("meta", hl))
    c = _contact_line(d)
    if c:
        out.append(("contact", c))

    summary = _s(d.get("summary") or d.get("professional_summary"))
    if summary:
        out += [("h2", "PROFESSIONAL SUMMARY"), ("p", summary)]

    skills = [x for x in (_s(y) for y in _list(d.get("skills") or d.get("skills_top"))) if x]
    if skills:
        out += [("h2", "CORE SKILLS"), ("p", " - ".join(skills[:24]))]

    exp = [e for e in _list(d.get("experience")) if isinstance(e, dict)]
    if exp:
        out.append(("h2", "PROFESSIONAL EXPERIENCE"))
        for e in exp:
            head = _join([e.get("title") or e.get("role"), e.get("company")], " - ")
            if head:
                out.append(("h3", head))
            dates = _s(e.get("dates")) or _join([e.get("start"), e.get("end")], " - ")
            meta = _join([dates, e.get("location")], " | ")
            if meta:
                out.append(("meta", meta))
            lede = _s(e.get("summary"))
            if lede:
                out.append(("p", lede))
            for b in _list(e.get("bullets") or e.get("highlights")):
                t = _s(b)
                if t:
                    out.append(("li", t))
            out.append(("spacer", ""))

    proj = [p for p in _list(d.get("projects")) if isinstance(p, dict)]
    if proj:
        out.append(("h2", "SELECTED PROJECTS"))
        for p in proj:
            nm = _s(p.get("name") or p.get("title"))
            if nm:
                out.append(("h3", nm))
            for b in _list(p.get("bullets") or p.get("highlights")):
                if _s(b):
                    out.append(("li", _s(b)))

    edu = [e for e in _list(d.get("education")) if isinstance(e, dict)]
    if edu:
        out.append(("h2", "EDUCATION"))
        for e in edu:
            head = _join([e.get("credential") or e.get("degree"), e.get("field_of_study")], ", ")
            if head:
                out.append(("h3", head))
            meta = _join([e.get("org") or e.get("school"),
                          e.get("years") or _join([e.get("from_year"), e.get("to_year")], "-"),
                          e.get("location")], " | ")
            if meta:
                out.append(("meta", meta))

    certs = [x for x in (_s(y) for y in _list(d.get("certifications"))) if x]
    if certs:
        out.append(("h2", "CERTIFICATIONS"))
        out += [("li", c) for c in certs]

    langs = [x for x in (_s(y) for y in _list(d.get("languages"))) if x]
    if langs:
        out += [("h2", "LANGUAGES"), ("p", " - ".join(langs))]

    return out


def _cover_blocks(d: dict) -> list:
    out = [("h1", _name_of(d))]
    hl = _headline_of(d)
    if hl:
        out.append(("meta", hl))
    c = _contact_line(d)
    if c:
        out.append(("contact", c))

    out.append(("meta", _s(d.get("date")) or _dt.date.today().strftime("%d %B %Y")))
    to = _join([d.get("company"), d.get("job_title")], " - ")
    if to:
        out.append(("p", to))
    out.append(("spacer", ""))

    sal = _s(d.get("salutation")) or "Dear Hiring Team"
    if not sal.lower().startswith("dear"):
        sal = "Dear " + sal
    out.append(("p", sal.rstrip(",") + ","))

    paras = [_s(p) for p in _list(d.get("paragraphs"))]
    if not any(paras):
        paras = [_s(d.get(k)) for k in ("opening", "proof", "fit", "close")]
    for p in paras:
        if p:
            out.append(("p", p))

    out += [("spacer", ""), ("p", _s(d.get("closing")) or "Sincerely,"), ("p", _name_of(d))]
    return out


def struct_to_blocks(d: dict, kind: str = "resume") -> list:
    """Turn the tailored-content dict into blocks. Missing fields are simply skipped."""
    d = d if isinstance(d, dict) else {}
    return _cover_blocks(d) if kind == "cover" else _resume_blocks(d)


# --------------------------------------------------------------------------- markdown -> blocks
def _inline(s: str) -> str:
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)
    s = re.sub(r"__(.+?)__", r"\1", s)
    s = re.sub(r"`(.+?)`", r"\1", s)
    s = re.sub(r"\[(.+?)\]\((.+?)\)", r"\1 (\2)", s)
    return _s(s)


def markdown_to_blocks(md: str, kind: str = "resume") -> list:
    """Accept the LLM's markdown when it hands us prose instead of JSON.

    Supported: '# H1', '## H2', '### H3', '- '/'* ' bullets, '**bold**' flattened to text,
    horizontal rules -> spacer. Anything else becomes a paragraph. Never raises.
    """
    out = []
    first_h1_seen = False
    for raw in _s(md).split("\n"):
        ln = raw.strip()
        if not ln:
            continue
        if ln.startswith("### "):
            out.append(("h3", _inline(ln[4:])))
        elif ln.startswith("## "):
            out.append(("h2", _inline(ln[3:]).upper()))
        elif ln.startswith("# "):
            txt = _inline(ln[2:])
            out.append(("h1", txt) if not first_h1_seen else ("h2", txt.upper()))
            first_h1_seen = True
        elif re.match(r"^[-*]\s+", ln):
            out.append(("li", _inline(re.sub(r"^[-*]\s+", "", ln))))
        elif re.match(r"^\d+[.)]\s+", ln):
            out.append(("li", _inline(re.sub(r"^\d+[.)]\s+", "", ln))))
        elif re.match(r"^([*_-]\s*){3,}$", ln):
            out.append(("spacer", ""))
        else:
            body = _inline(ln)
            if len(out) <= 2 and ("@" in body or re.search(r"\+\d", body)) and len(body) < 200:
                out.append(("contact", body))
            else:
                out.append(("p", body))
    return out or [("p", "")]


def to_blocks(md_or_struct: Any, kind: str = "resume") -> list:
    """The one entry point both writers use."""
    if isinstance(md_or_struct, dict):
        return struct_to_blocks(md_or_struct, kind)
    if isinstance(md_or_struct, (list, tuple)) and md_or_struct and \
            isinstance(md_or_struct[0], (list, tuple)) and len(md_or_struct[0]) == 2:
        return [(_s(a), _s(b)) for a, b in md_or_struct]
    return markdown_to_blocks(_s(md_or_struct), kind)


# --------------------------------------------------------------------------- DOCX writer
def _docx_base():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = FONT_DOCX
    st.font.size = Pt(10.5)
    pf = st.paragraph_format
    pf.space_after = Pt(4)
    pf.line_spacing = 1.08
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Pt(40)
        s.left_margin = s.right_margin = Pt(48)
    return doc


def _write_docx(blocks: list, path: str) -> str:
    doc = _docx_base()
    for kind, text in blocks:
        text = _s(text)
        if kind == "spacer":
            doc.add_paragraph("")
            continue
        if not text:
            continue
        if kind == "h1":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(20)
            r.font.color.rgb = ACCENT
            p.paragraph_format.space_after = Pt(2)
        elif kind == "contact":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(10)
        elif kind == "meta":
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.italic = True
            r.font.size = Pt(9.5)
            p.paragraph_format.space_after = Pt(3)
        elif kind == "h2":
            p = doc.add_paragraph(style="Heading 1")
            r = p.add_run(text.upper())
            r.bold = True
            r.font.size = Pt(12)
            r.font.name = FONT_DOCX
            r.font.color.rgb = ACCENT
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
        elif kind == "h3":
            p = doc.add_paragraph(style="Heading 2")
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = FONT_DOCX
            r.font.color.rgb = RGBColor(0, 0, 0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
        elif kind == "li":
            p = doc.add_paragraph(text, style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _ensure_dir(path)
    doc.save(path)
    return path


# --------------------------------------------------------------------------- PDF writer
def _esc(s: str) -> str:
    return _s(s).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


_PDF_STYLES = {
    "h1": ParagraphStyle("h1", fontName=FONT_PDF_B, fontSize=19, leading=22,
                         textColor="#1F364D", spaceAfter=2, alignment=TA_LEFT),
    "contact": ParagraphStyle("contact", fontName=FONT_PDF, fontSize=9, leading=12,
                              textColor="#333333", spaceAfter=9),
    "meta": ParagraphStyle("meta", fontName=FONT_PDF, fontSize=9, leading=12,
                           textColor="#444444", spaceAfter=3),
    "h2": ParagraphStyle("h2", fontName=FONT_PDF_B, fontSize=11.5, leading=14,
                         textColor="#1F364D", spaceBefore=10, spaceAfter=3),
    "h3": ParagraphStyle("h3", fontName=FONT_PDF_B, fontSize=10.5, leading=13,
                         spaceBefore=6, spaceAfter=1),
    "p": ParagraphStyle("p", fontName=FONT_PDF, fontSize=10, leading=13.5, spaceAfter=4),
    "li": ParagraphStyle("li", fontName=FONT_PDF, fontSize=10, leading=13.5, spaceAfter=2),
}


def _ensure_dir(path: str) -> None:
    d = os.path.dirname(os.path.abspath(path))
    if d:
        os.makedirs(d, exist_ok=True)


def _write_pdf(blocks: list, path: str, title: str = "") -> str:
    _ensure_dir(path)
    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=17 * mm, rightMargin=17 * mm,
                            topMargin=15 * mm, bottomMargin=15 * mm,
                            title=_s(title) or "Document", author=_s(title))
    flow = []
    pending = []

    def flush():
        if pending:
            flow.append(ListFlowable(
                [ListItem(Paragraph(_esc(t), _PDF_STYLES["li"]), leftIndent=10)
                 for t in pending],
                bulletType="bullet", start="bulletchar", leftIndent=12,
                bulletFontName=FONT_PDF, bulletFontSize=9))
            del pending[:]

    for kind, text in blocks:
        text = _s(text)
        if kind == "li":
            if text:
                pending.append(text)
            continue
        flush()
        if kind == "spacer":
            flow.append(Spacer(1, 5))
            continue
        if not text:
            continue
        flow.append(Paragraph(_esc(text.upper() if kind == "h2" else text),
                              _PDF_STYLES.get(kind, _PDF_STYLES["p"])))
    flush()
    if not flow:
        flow = [Paragraph("", _PDF_STYLES["p"])]
    doc.build(flow)
    return path


def _title_of(blocks: list) -> str:
    for k, t in blocks:
        if k == "h1":
            return _s(t)
    return "Document"


# --------------------------------------------------------------------------- public API
def resume_docx(md_or_struct: Any, path: str) -> str:
    """Write an ATS-safe resume .docx. Accepts a struct dict or markdown text."""
    return _write_docx(to_blocks(md_or_struct, "resume"), path)


def cover_docx(md_or_struct: Any, path: str) -> str:
    """Write a cover-letter .docx. Accepts a struct dict or markdown text."""
    return _write_docx(to_blocks(md_or_struct, "cover"), path)


def resume_pdf(md_or_struct: Any, path: str) -> str:
    b = to_blocks(md_or_struct, "resume")
    return _write_pdf(b, path, title=_title_of(b))


def cover_pdf(md_or_struct: Any, path: str) -> str:
    b = to_blocks(md_or_struct, "cover")
    return _write_pdf(b, path, title=_title_of(b))


def to_pdf(md_or_struct: Any, path: str, kind: str = "resume") -> str:
    """Render to PDF with reportlab (pure Python - no system binaries required)."""
    b = to_blocks(md_or_struct, kind)
    return _write_pdf(b, path, title=_title_of(b))


def write_all(resume: Any, cover: Any, outdir: str) -> dict:
    """Write resume.docx/.pdf + cover_letter.docx/.pdf into outdir.

    A failure in one renderer never blocks the others - the caller gets whatever
    succeeded plus an 'errors' map, instead of an exception and zero files.
    """
    os.makedirs(outdir, exist_ok=True)
    jobs = [
        ("resume.docx", resume_docx, resume),
        ("resume.pdf", resume_pdf, resume),
        ("cover_letter.docx", cover_docx, cover),
        ("cover_letter.pdf", cover_pdf, cover),
    ]
    files = {}
    errors = {}
    for name, fn, payload in jobs:
        if payload is None:
            continue
        p = os.path.join(outdir, name)
        try:
            fn(payload, p)
            files[name] = p
        except Exception as e:
            errors[name] = "%s: %s" % (type(e).__name__, e)
    return {"files": files, "errors": errors}
